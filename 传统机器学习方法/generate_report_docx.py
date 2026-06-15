"""生成个人工作报告 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

# ============ 封面 ============
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CROHME 数学公式识别\n传统机器学习路径实验报告')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('日期：2026年5月19日\n').font.size = Pt(12)
info.add_run('环境：Python 3.12 / scikit-learn / OpenCV\n').font.size = Pt(12)
info.add_run('数据：CROHME 2011-2014 手写数学公式数据集').font.size = Pt(12)

doc.add_page_break()

# ============ 第一章 项目概述 ============
doc.add_heading('一、项目背景与目标', level=1)
doc.add_paragraph(
    '本项目基于 CROHME 手写数学公式数据集，探索传统计算机视觉 + 机器学习方法在数学公式识别任务上的可行性，'
    '为后续深度学习方案建立性能基线。'
)
doc.add_paragraph('技术路线：InkML 解析 → 笔画渲染 → 数据增强 → HOG+LBP 特征提取 → StandardScaler → PCA 降维 → SVM 分类')

# ============ 第二章 踩坑全记录 ============
doc.add_heading('二、实验过程中的问题与解决方案', level=1)

doc.add_heading('2.1 第一次运行：数据加载与标注格式问题', level=2)
doc.add_paragraph(
    '最初尝试直接使用 CROHME_test_2011 目录的数据，发现 load_from_folder 返回 0 条样本。'
    '排查发现该目录下的 inkml 文件 annotation 标签中 type 字段为 "UI" 而非 "truth"，'
    '即测试集不包含 ground truth 标签。'
)
p = doc.add_paragraph()
p.add_run('解决方案：').bold = True
p.add_run('改用全量 crohme_data 目录遍历，利用 os.walk 递归搜索所有含 truth 标签的 inkml 文件（2012/2013/2014 年训练数据），成功加载 22,021 条样本。')

doc.add_heading('2.2 第二次运行：549 类 GridSearchCV 时间爆炸', level=2)
doc.add_paragraph(
    '首次全量运行传统模型时，将全部 549 个公式类别直接送入 GridSearchCV，参数网格包含 linear(4 种 C) + rbf(20 种 C×gamma) + poly(6 种) '
    '共 30 种组合，配合 5-fold 交叉验证。SVM 使用 OVR 策略，每个拟合需训练 549 个二分类器。'
    '预估单特征配置需运行数小时，3 组特征配置完全无法在合理时间内完成。'
)
p = doc.add_paragraph()
p.add_run('解决方案：').bold = True
p.add_run(
    '（1）改用 Top-N 策略，仅保留样本数最多的公式类别（min_samples >= 25 → 154 类，'
    '最终 Top-30 → 3,937 样本）；（2）将完整 GridSearchCV 替换为 RandomizedSearchCV(n_iter=8)，'
    '搜索量从 30 种组合降至 8 种随机采样；（3）3-fold 替代 5-fold，搜索量再降 40%。'
)

doc.add_heading('2.3 第三次运行：HalvingGridSearchCV 实验导入失败', level=2)
doc.add_paragraph(
    '尝试引入 HalvingGridSearchCV 利用 successive halving 策略加速：先用少量数据淘汰劣质参数，仅对优秀候选执行完整评估。'
    '但 sklearn 将 HalvingGridSearchCV 标记为 experimental API，直接导入触发 ImportError。'
)
p = doc.add_paragraph()
p.add_run('解决方案：').bold = True
p.add_run('添加 from sklearn.experimental import enable_halving_search_cv 前置导入语句。但后续实测发现 Windows 环境下 n_jobs=-1 多进程开销极大，反而比串行更慢，最终弃用 HalvingGridSearchCV，回归 RandomizedSearchCV。')

doc.add_heading('2.4 第四次运行：n_jobs=-1 导致进程崩溃', level=2)
doc.add_paragraph(
    '设置 n_jobs=-1 利用全部 CPU 核心并行训练后，监控发现同时产生 18+ 个 Python 子进程。'
    '任务被手动停止后，linlinear 子进程未被正确清理，成为僵尸进程继续占用内存。'
    '后续多次启动/停止累积了大量僵尸进程，导致某次运行时系统内存耗尽，进程被 Windows 强制 kill（无 Python traceback）。'
)
p = doc.add_paragraph()
p.add_run('解决方案：').bold = True
p.add_run(
    '（1）将 n_jobs 从 -1 改为 1，强制串行执行避免多进程开销；'
    '（2）每次停止任务后执行 Get-Process python* | Stop-Process -Force 彻底清理；'
    '（3）添加 warnings.filterwarnings("ignore") 抑制 liblinear 收敛警告减少日志噪音。'
)

doc.add_heading('2.5 第五次运行：配置 2 崩溃分析', level=2)
doc.add_paragraph(
    '实验设计了 3 组特征配置对比：配置 1 (HOG 9方向, ppc=8)、配置 2 (HOG 9方向, ppc=4 + LBP 16, R=2)、'
    '配置 3 (HOG 18方向, ppc=8 + LBP 默认模式)。配置 1 在 34 分钟内完成（910维 → PCA 303维），'
    '但配置 2（4,374维 → PCA 1,071维）的 LinearSVC 训练 86+ 分钟后进程消失。'
)
p = doc.add_paragraph()
p.add_run('根因分析：').bold = True
doc.add_paragraph(
    '配置 2 的 PCA 后维度 1,071（是配置 1 的 3.5 倍），每个 OVR 二分类器训练时间与特征维度成正比。'
    '8轮 × 3折 × 30类 = 720 个二分类器，单次约 200s，预计总耗时 80+ 分钟。'
    '在长时间高内存占用下，残留僵尸进程导致系统 OOM，进程被强制终止。'
)
p = doc.add_paragraph()
p.add_run('解决方案：').bold = True
p.add_run(
    '（1）移除 ppc=4 的细粒度 HOG 配置（性价比极低）；（2）移除 default LBP 模式（256 bins 稀疏直方图对 48×48 小图无效）；'
    '（3）仅保留 HOG(9,8,2)+LBP(8,1,uniform) 单配置——该配置是文献中最标准、最稳定的参数组合。'
)

doc.add_heading('2.6 IMG_SIZE 优化：64 → 48', level=2)
doc.add_paragraph(
    '最初使用 64×64 图片（HOG 1,764维 → PCA 541维）。为加速训练，将图片缩小至 48×48，'
    'HOG 降至 900维，PCA 仅需 303 维即可保留 95% 方差。特征维度减半，训练时间约减少 40%，而准确率无明显下降。'
)

p = doc.add_paragraph()
p.add_run('尺寸对比：').bold = True
table = doc.add_table(rows=4, cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['IMG_SIZE', 'HOG维', 'LBP维', '总维', 'PCA后', '预计耗时']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ['64×64', '1,764', '10', '1,774', '541', '~60min'],
    ['48×48', '900', '10', '910', '303', '~34min'],
    ['缩减', '-49%', '0', '-49%', '-44%', '-43%'],
]
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r+1].cells[c].text = val

doc.add_heading('2.7 代码架构演变', level=2)

items = [
    ('preprocess.py v1.0', '基础 InkML 解析 + 简单渲染，无增强，无过滤'),
    ('preprocess.py v2.0', '增加自适应居中渲染、数据增强 pipeline、类别过滤、分布统计'),
    ('traditional_model.py v1.0', 'HOG+LBP+PCA+SVM 基础版，固定参数 C=10，无标准化'),
    ('traditional_model.py v1.1', '+ StandardScaler 标准化、GridSearchCV 调参、多核对比'),
    ('traditional_model.py v2.0', '+ 实验记录 JSON、分类报告、混淆矩阵、误分类分析'),
    ('traditional_model.py v2.1', '+ 12 张全流程可视化（样本图库→增强对比→HOG/LBP→PCA→混淆→仪表盘）'),
    ('traditional_model.py v2.2', 'GridSearchCV → RandomizedSearchCV、n_jobs=-1 → 1、IMG_SIZE 64 → 48'),
    ('最终版本', '配置 1 单跑 + Top-30 类 + 48×48 + RandomizedSearchCV + 全可视化 + JSON 报告'),
]
for title, desc in items:
    p = doc.add_paragraph()
    p.add_run(f'{title}：').bold = True
    p.add_run(desc)

doc.add_page_break()

# ============ 第三章 技术方案 ============
doc.add_heading('三、技术方案详述', level=1)

doc.add_heading('3.1 数据预处理', level=2)
doc.add_paragraph('数据来源：CROHME 2011-2014 年含 truth 标注的训练数据，通过 os.walk 递归扫描全部 inkml 文件。')
doc.add_paragraph('标签过滤：取 Counter(y).most_common(30)，仅保留样本数最多的 30 个公式类（每类 124-135 个样本）。')

doc.add_heading('3.2 图像渲染', level=2)
doc.add_paragraph('render_strokes() 将笔画轨迹渲染为 48×48 灰度图：')
doc.add_paragraph('• 以笔画质心为基准自适应缩放，保持宽高比')
doc.add_paragraph('• 抗锯齿线条 (cv2.LINE_AA)，线宽 2px')
doc.add_paragraph('• 空白边缘 10px padding，防止笔画贴边')

doc.add_heading('3.3 数据增强', level=2)
doc.add_paragraph('augment_dataset() 对每类样本做 on-the-fly 增强，确保各类均衡：')
doc.add_paragraph('• 旋转 ±10°、缩放 ±15%、平移 ±2px')
doc.add_paragraph('• 弹性形变（50% 概率），模拟不同书写风格')
doc.add_paragraph('• 目标：每类 ≥60 样本，factor=2x 扩展')

doc.add_heading('3.4 特征工程', level=2)
p = doc.add_paragraph()
p.add_run('HOG (Histogram of Oriented Gradients)：').bold = True
p.add_run('orientations=9, pixels_per_cell=(8,8), cells_per_block=(2,2) → 900维。用于捕捉笔画边缘方向与轮廓信息。')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('LBP (Local Binary Pattern)：').bold = True
p.add_run('P=8, R=1, method="uniform" → 10 bins 直方图。用于捕捉局部纹理模式（笔画交叉点、拐角）。')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('特征标准化：').bold = True
p.add_run('StandardScaler 将 HOG+LBP 拼接特征标准化为均值0、方差1，确保 PCA 不被量纲差异主导。')

doc.add_heading('3.5 PCA 降维', level=2)
doc.add_paragraph('PCA(n_components=0.95) 保留 95% 累积方差，910维 → 303维，压缩率 67%。'
                  '仅在训练集上拟合 PCA，测试集使用相同变换矩阵。')

doc.add_heading('3.6 SVM 训练与调参', level=2)
doc.add_paragraph('使用 RandomizedSearchCV 对 LinearSVC 和 SVC(RBF) 分别搜索最优参数：')
doc.add_paragraph('• LinearSVC：C ∈ [10⁻³, 10³]，对数均匀采样 8 轮 → 最优 C = 0.039')
doc.add_paragraph('• RBF：C ∈ [10⁻², 10³]，γ ∈ [10⁻³, 10¹]，采样 8 轮 → 最优 C=0.356, γ=0.001')
doc.add_paragraph('• 交叉验证：3-fold StratifiedKFold')
doc.add_paragraph('• 最终模型：LinearSVC(C=0.039, class_weight="balanced")')

# ============ 第四章 实验结果 ============
doc.add_heading('四、实验结果', level=1)

doc.add_heading('4.1 最终实验结果（Top-30 类）', level=2)

table2 = doc.add_table(rows=8, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
data2 = [
    ('指标', '数值'),
    ('数据集', 'CROHME Top-30，7,874样本'),
    ('训练/测试', '5,511 / 2,363 (70/30)'),
    ('特征', 'HOG(9,8,2)+LBP(8,1,uniform)'),
    ('特征维度', '910 → PCA 303（95%方差）'),
    ('最优模型', 'LinearSVC, C=0.039'),
    ('CV Accuracy', '66.99% ± 0.18%'),
    ('Test Accuracy', '68.73%'),
]
for i, (k, v) in enumerate(data2):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v
    if i == 0:
        for c in range(2):
            table2.rows[0].cells[c].paragraphs[0].runs[0].font.bold = True

doc.add_heading('4.2 核函数对比', level=2)
table3 = doc.add_table(rows=3, cols=4)
table3.style = 'Light Grid Accent 1'
for i, h in enumerate(['核函数', '最优参数', 'CV Accuracy', '训练耗时']):
    table3.rows[0].cells[i].text = h
table3.rows[1].cells[0].text = 'LinearSVC'
table3.rows[1].cells[1].text = 'C=0.039'
table3.rows[1].cells[2].text = '66.99%'
table3.rows[1].cells[3].text = '1,973s (32.9min)'
table3.rows[2].cells[0].text = 'RBF'
table3.rows[2].cells[1].text = 'C=0.356, γ=0.001'
table3.rows[2].cells[2].text = '66.67%'
table3.rows[2].cells[3].text = '56s'

doc.add_heading('4.3 不同规模实验结果', level=2)
table4 = doc.add_table(rows=4, cols=5)
table4.style = 'Light Grid Accent 1'
for i, h in enumerate(['实验', '类别数', '样本数', 'Test Acc', '随机基线']):
    table4.rows[0].cells[i].text = h
scale_data = [
    ('quick_test', '10', '2,700', '77.78%', '10.0%'),
    ('visualize_predictions', '15', '4,050', '72.92%', '6.7%'),
    ('traditional_model', '30', '7,874', '68.73%', '3.3%'),
]
for r, row in enumerate(scale_data):
    for c, val in enumerate(row):
        table4.rows[r+1].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph('规律：类别数与准确率呈负相关，每增加约 15 类，准确率下降 ~5%。')

doc.add_page_break()
# ============ 第五章 关键发现 ============
doc.add_heading('五、关键发现与分析', level=1)

doc.add_heading('5.1 LinearSVC 优于 RBF', level=2)
doc.add_paragraph(
    '在小样本多类别场景下（30 类 × ~180 训练样本/类），线性核的泛化性优于 RBF 核。'
    'PCA 已将特征映射至近线性可分空间，RBF 的非线性变换反而增加过拟合风险。'
    '此外，RBF 参数空间（C × γ）更复杂，RandomizedSearchCV 在有限的 8 次采样中难以命中最优组合。'
)

doc.add_heading('5.2 PCA 降维效果显著', level=2)
doc.add_paragraph(
    '910 维 → 303 维，压缩率 67%，信息损失仅 5%。说明 HOG 特征在相邻 cell 间高度相关，存在大量冗余。'
    '48×48 图片 + ppc=8 单元比是性价比最优配置：比 64×64 节省 49% 特征量，准确率无实质下降。'
)

doc.add_heading('5.3 数据增强是必需步骤', level=2)
doc.add_paragraph(
    '不做增强时每类仅 ~95 训练样本，30 个 OVR 分类器难以收敛。增强后每类 ~180 训练样本，CV 稳定在 67%。'
    '弹性形变（50% 概率）模拟了不同书写风格，是提升泛化性的关键增强操作。'
)

doc.add_heading('5.4 传统方法的根本局限', level=2)
doc.add_paragraph(
    'HOG+LBP 是纯空间域特征，完全丢失了手写笔画的时间顺序信息。两个视觉相似但书写顺序不同的公式，'
    '传统 CV 方法几乎无法区分。此外，完整 LaTeX 公式作为标签粒度太粗，实际应用需要符号级别的识别。'
)

# ============ 第六章 可视化图表 ============
doc.add_heading('六、可视化成果', level=1)
doc.add_paragraph('共生成 14 张实验可视化图表，存放于 experiments/figures/ 目录：')

figs = [
    ('01_sample_grid', 'Top-20 类样本图库'),
    ('02_augmentation_comparison', '数据增强前后对比'),
    ('03_class_distribution', '类别分布直方图'),
    ('04_hog_visualization', 'HOG 梯度特征可视化'),
    ('05_lbp_visualization', 'LBP 纹理特征可视化'),
    ('06_pca_variance', 'PCA 累计方差曲线（303维@95%）'),
    ('07_dimension_comparison', '原始维 vs PCA 维对比'),
    ('08_accuracy_comparison', 'CV vs Test 准确率对比'),
    ('09_kernel_time', '核函数选择与训练耗时'),
    ('10_confusion_heatmap', '混淆矩阵热力图 (Top-15)'),
    ('11_misclassification', 'Top 误分类对分析'),
    ('12_summary_dashboard', '综合实验仪表盘'),
    ('13_prediction_samples', '预测结果展示（原图+真实+预测）'),
    ('14_per_class_accuracy', '各类别准确率横向对比'),
]
for name, desc in figs:
    doc.add_paragraph(f'• {name}.png：{desc}')

# ============ 第七章 后续方向 ============
doc.add_heading('七、后续改进方向', level=1)

improvements = [
    ('在线时序特征', '加入笔画方向直方图、书写速度/加速度、曲率等时序特征，利用手写数据的时间维度', '预期 +3~5%'),
    ('符号级识别', '使用 CROHME 2013/2014 的 isolated symbol 数据集或引入符号分割模型，将任务从公式级降低到符号级', '预期 +10~15%'),
    ('多特征融合', '在线特征(时序) + 离线特征(图像) 拼接，Early Fusion 或 Late Fusion', '预期 +5~8%'),
    ('深度学习', 'CNN 提取图像特征 + RNN/LSTM 处理时序 + Attention 机制，端到端训练', '预期 +15~20%'),
]
for title, desc, expected in improvements:
    p = doc.add_paragraph()
    p.add_run(f'{title}：').bold = True
    p.add_run(f'{desc}（{expected}）')

# ============ 第八章 总结 ============
doc.add_heading('八、总结', level=1)
doc.add_paragraph(
    '本项目从零搭建了 CROHME 手写数学公式识别的传统 CV+ML 完整 pipeline，经历了从数据加载失败、'
    'GridSearchCV 时间爆炸、多进程崩溃、高维特征 OOM、到最终稳定收敛的完整踩坑过程。在此期间：'
)
doc.add_paragraph('• 实现了 InkML 解析与笔画渲染、5 种数据增强、HOG+LBP 双特征融合')
doc.add_paragraph('• 优化了 StandardScaler+PCA 特征工程管线，特征维度从 1,774 降至 303')
doc.add_paragraph('• 从 GridSearchCV → HalvingGridSearchCV → RandomizedSearchCV，最终选定高效调参策略')
doc.add_paragraph('• 解决了 Windows 多进程管理、Python 僵尸进程清理、收敛警告抑制等工程问题')
doc.add_paragraph('• 12 张全流程可视化图表 + JSON 结构化实验记录，确保可复现性')
doc.add_paragraph()
doc.add_paragraph(
    '最终模型 HOG(9,8,2)+LBP(8,1,uniform) → PCA(303维) → LinearSVC(C=0.039) 在 CROHME Top-30 公式识别任务上 '
    '达到 68.73% 测试准确率，超出随机基线（3.3%）20 倍，验证了传统 CV 方法在数学公式识别上的可行性。'
    '但距实用水平（>90%）仍有差距，根本瓶颈在于传统特征无法捕获手写的时序特性。'
    '建议下一步引入在线笔画特征或迁移至深度学习方法。'
)

# ============ 附录 ============
doc.add_heading('附录：产出物清单', level=1)
artifacts = [
    ('preprocess.py', 'InkML 解析、笔画渲染、数据增强 pipeline'),
    ('traditional_model.py', '完整 HOG+LBP+PCA+SVM 训练 pipeline + 12 图可视化'),
    ('quick_test.py', '快速验证脚本（Top-10 类，8 分钟出结果）'),
    ('demo_one_sample.py', '单样本演示：原始数据 → 图像 → 符号拆分'),
    ('visualize_predictions.py', '预测效果可视化：原图 + 真实公式 + 预测公式'),
    ('finish_report.py', '补全实验报告与缺失图表'),
    ('experiments/report_*.json', '结构化实验数据（可复现）'),
    ('experiments/figures/01-14_*.png', '14 张全流程可视化图表'),
    ('experiments/个人工作报告.md', 'Markdown 版完整报告'),
    ('experiments/个人工作报告.docx', '本 Word 文档'),
]
for name, desc in artifacts:
    doc.add_paragraph(f'• {name}：{desc}')

# ============ 保存 ============
output_path = r'd:\Users\admin\Desktop\手写数字识别\experiments\个人工作报告.docx'
doc.save(output_path)
print(f'Word 文档已保存：{output_path}')
